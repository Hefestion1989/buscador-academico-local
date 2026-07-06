from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class TextSegment:
    text: str
    location: str


def extract_segments(path: Path) -> list[TextSegment]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return extract_pdf(path)
    if extension == ".docx":
        return extract_docx(path)
    if extension in {".txt", ".md"}:
        return extract_txt(path)
    if extension == ".rtf":
        return extract_rtf(path)
    return []


def extract_pdf(path: Path) -> list[TextSegment]:
    try:
        return extract_pdf_with_pymupdf(path)
    except Exception:
        return extract_pdf_with_pypdf(path)


def extract_pdf_with_pymupdf(path: Path) -> list[TextSegment]:
    import fitz

    segments: list[TextSegment] = []
    with fitz.open(path) as document:
        for index, page in enumerate(document, start=1):
            text = clean_text(page.get_text("text") or "")
            if text:
                segments.append(TextSegment(text=text, location=f"pagina {index}"))
    return segments


def extract_pdf_with_pypdf(path: Path) -> list[TextSegment]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    segments: list[TextSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = clean_text(text)
        if text:
            segments.append(TextSegment(text=text, location=f"pagina {index}"))
    return segments


def extract_docx(path: Path) -> list[TextSegment]:
    from docx import Document

    document = Document(str(path))
    segments: list[TextSegment] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = clean_text(paragraph.text)
        if text:
            segments.append(TextSegment(text=text, location=f"parrafo {index}"))

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            segments.append(
                TextSegment(
                    text="\n".join(rows),
                    location=f"tabla {table_index}",
                )
            )

    return segments


def extract_txt(path: Path) -> list[TextSegment]:
    text = read_text_with_fallback(path)
    segments: list[TextSegment] = []
    line_number = 1

    for block in text.split("\n\n"):
        cleaned = clean_text(block)
        if cleaned:
            segments.append(TextSegment(text=cleaned, location=f"linea {line_number}"))
        line_number += block.count("\n") + 2

    return segments


def extract_rtf(path: Path) -> list[TextSegment]:
    text = read_text_with_fallback(path)
    text = strip_basic_rtf(text)
    cleaned = clean_text(text)
    if not cleaned:
        return []
    return [TextSegment(text=cleaned, location="contenido rtf")]


def strip_basic_rtf(text: str) -> str:
    import re

    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("\\{", "{").replace("\\}", "}").replace("\\\\", "\\")
    text = re.sub(r"[{}]", " ", text)
    return text


def read_text_with_fallback(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def clean_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
